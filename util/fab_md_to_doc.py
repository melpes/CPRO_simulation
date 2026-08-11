"""FAB 방문자료 md 를 docx 로 변환하고 서식을 다듬는다.

pandoc 이 만든 docx 는 한글 폰트가 잡히지 않아 style 을 직접 손본다.
"""
import subprocess, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ── 하드코딩 ────────────────────────────────────────────────────────────
BASE = r"C:/Users/KangTaehui/KG/keti/CPRO_조립공정/시뮬레이션/Package/docs/FAB"
NAME = "핵심공정선정_방문자료"
MD = f"{BASE}/{NAME}.md"
# 인자로 파일명을 주면 그 이름으로 저장한다 (기존 docx 가 열려 있을 때 사용)
DOCX = f"{BASE}/{sys.argv[1]}.docx" if len(sys.argv) > 1 else f"{BASE}/{NAME}.docx"

FONT = "맑은 고딕"
FONT_MONO = "D2Coding"          # 없으면 아래에서 Consolas 로 대체
BODY_PT = 10
TABLE_PT = 9


def set_font(run, name, size=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


# ── 1. pandoc 변환 ──────────────────────────────────────────────────────
subprocess.run(["pandoc", MD, "-o", DOCX, "--from=gfm", "--toc", "--toc-depth=2"],
               check=True)

# ── 2. 서식 다듬기 ──────────────────────────────────────────────────────
doc = Document(DOCX)

# 기본 스타일
st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(BODY_PT)
st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
st.paragraph_format.space_after = Pt(4)
st.paragraph_format.line_spacing = 1.3

# 제목 — 번호 없는 굵은 검정
for lv, sz in ((1, 16), (2, 13), (3, 11), (4, 10)):
    try:
        h = doc.styles[f"Heading {lv}"]
    except KeyError:
        continue
    h.font.name = FONT
    h.font.size = Pt(sz)
    h.font.bold = True
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    h.paragraph_format.space_before = Pt(14 if lv <= 2 else 10)
    h.paragraph_format.space_after = Pt(6)

# 여백
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.0)
    s.left_margin = s.right_margin = Cm(2.0)

def add_borders(table):
    """pandoc 이 만든 docx 엔 Table Grid 스타일이 없어 테두리를 직접 넣는다."""
    from docx.oxml import OxmlElement
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "999999")
        borders.append(e)
    tblPr.append(borders)


# 표 — 격자 + 헤더 굵게 + 글자 축소
for t in doc.tables:
    add_borders(t)
    t.autofit = True
    for i, row in enumerate(t.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.1
                for r in p.runs:
                    set_font(r, FONT, TABLE_PT, bold=(i == 0) or r.font.bold)

# 코드 블록 — 고정폭
for p in doc.paragraphs:
    if p.style.name in ("Source Code", "Verbatim Char", "HTML Preformatted"):
        for r in p.runs:
            set_font(r, "Consolas", 9)

doc.save(DOCX)
print("docx :", DOCX)

# ── 3. hwp 저장 ─────────────────────────────────────────────────────────
print(f"       표 {len(doc.tables)}개 · 문단 {len(doc.paragraphs)}개")
