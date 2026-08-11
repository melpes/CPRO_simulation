"""v2 docx 를 복사한 뒤 지정한 절만 갈아끼운다.

md 에서 통째로 재생성하면 사용자가 docx 에서 맞춘 서식이 날아가므로,
문단·표 단위로 해당 구간만 교체한다.
"""
import shutil, copy
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = r"C:/Users/KangTaehui/KG/keti/CPRO_조립공정/시뮬레이션/Package/docs/FAB"
SRC = f"{BASE}/핵심공정선정_방문자료_v2.docx"
DST = f"{BASE}/핵심공정선정_방문자료_v3.docx"

shutil.copyfile(SRC, DST)
doc = Document(DST)
body = doc.element.body


# ── 헬퍼 ────────────────────────────────────────────────────────────────
def heading_el(text, level=None, exact=False):
    """제목 문단의 XML 요소를 찾는다.

    exact=True 면 완전 일치. "용어" 가 "스크린 프린팅 용어" 에 걸리는 것을 막는다.
    """
    hits = []
    for p in doc.paragraphs:
        if not p.style.name.startswith("Heading"):
            continue
        if level and not p.style.name.endswith(str(level)):
            continue
        t = p.text.strip()
        if (t == text) if exact else (text in t):
            hits.append(p._element)
    if not hits:
        raise KeyError(text)
    if exact and len(hits) > 1:
        raise KeyError(f"{text} — 제목이 {len(hits)}개 중복")
    return hits[0]


def section_range(start_el, end_el):
    """start_el 과 end_el(둘 다 제목) 사이의 요소 목록. 끝을 자동 판정하지 않는다."""
    kids = list(body)
    i, j = kids.index(start_el), kids.index(end_el)
    if j <= i:
        raise ValueError("끝 제목이 시작 제목보다 앞에 있음")
    return kids[i + 1:j]


def borders(tbl):
    b = OxmlElement("w:tblBorders")
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        x = OxmlElement(f"w:{e}")
        x.set(qn("w:val"), "single")
        x.set(qn("w:sz"), "4")
        x.set(qn("w:color"), "999999")
        b.append(x)
    tbl._tbl.tblPr.append(b)


def style_run(r, size=None, bold=None):
    r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if size:
        r.font.size = Pt(size)
    if bold is not None:
        r.font.bold = bold


def make_para(text, style="Normal", size=10):
    p = doc.add_paragraph(style=style)
    for chunk, bold in parse_bold(text):
        r = p.add_run(chunk)
        style_run(r, size, bold)
    p.paragraph_format.space_after = Pt(4)
    return p._element


def parse_bold(text):
    """**굵게** 표기를 (텍스트, 굵기) 로 쪼갠다."""
    out, buf, bold = [], "", False
    i = 0
    while i < len(text):
        if text[i:i + 2] == "**":
            if buf:
                out.append((buf, bold))
                buf = ""
            bold = not bold
            i += 2
        else:
            buf += text[i]
            i += 1
    if buf:
        out.append((buf, bold))
    return out or [(text, False)]


def make_table(rows, header=True):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    borders(t)
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            c = t.cell(ri, ci)
            c.text = ""
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            for chunk, bold in parse_bold(str(cell)):
                r = p.add_run(chunk)
                style_run(r, 9, bold or (header and ri == 0))
    return t._tbl


def replace_section(title, end_title, blocks, level=2):
    """title 과 end_title 사이를 blocks 로 교체. 끝을 명시해 과삭제를 막는다."""
    h = heading_el(title, level, exact=True)
    e = heading_el(end_title, exact=True)
    for el in section_range(h, e):
        el.getparent().remove(el)
    anchor = h
    for kind, data in blocks:
        el = make_para(data) if kind == "p" else make_table(data)
        anchor.addnext(el)
        anchor = el


def insert_before(title, blocks, level=3):
    """제목 앞에 blocks 를 끼워 넣는다."""
    h = heading_el(title, level)
    made = []
    for kind, data in blocks:
        made.append(make_para(data, style=f"Heading {level}") if kind == "h"
                    else make_para(data) if kind == "p" else make_table(data))
    for el in reversed(made):
        h.addprevious(el)


# ── ② 로그 실측 교체 ────────────────────────────────────────────────────
replace_section("로그 실측", "용어", [
    ("p", "로그로 알 수 있는 것은 **기판이 챔버를 얼마나 점유했는지**와 **설비가 언제 돌았는지** 두 가지다. "
          "공정 조건(파워·가스·시간)은 runsheet 에 있으므로, 로그는 그 조건이 실제로 얼마나 걸렸는지를 보여준다."),
    ("p", "**챔버 점유 시간 = 레시피 시간 + 약 92초**"),
    ("p", "레시피 이름에 공정 시간이 들어 있어(CF4_O2_600W_60s) 실제 점유 시간과 대조가 가능했다. "
          "Dry etcher 128런·PECVD 72런 전부에서 차이가 91~93초로 일정했다."),
    ("t", [["레시피", "레시피상 시간", "실제 점유", "차이"],
           ["Mo_S.D_250_500_20s", "20 s", "112 s", "+92"],
           ["CF4_O2_600W_60s", "60 s", "152 s", "+92"],
           ["Cl2_Ar_O2_220s", "220 s", "311 s", "+91"],
           ["MO_400W_360S", "360 s", "452 s", "+92"]]),
    ("p", "92초는 기판을 넣고 진공을 잡고 빼는 데 드는 고정 시간으로 보임. "
          "→ 시뮬의 설비 점유 시간은 **runsheet 공정시간 + 92초** 로 잡을 수 있음"),
    ("p", "**설비 가동 현황**"),
    ("t", [["항목", "값", "의미"],
           ["설비 가동률", "**4~7%**", "5개월 중 실제로 공정이 돈 시간의 비율"],
           ["런 간격", "중앙값 약 15분", "한 런이 끝나고 다음 런까지"],
           ["챔버 간 반송", "13~19초", "반송은 병목이 아님"],
           ["로드락", "약 4분", "대기압 ↔ 진공 전환"]]),
    ("p", "**병목은 설비가 아니라 작업자**"),
    ("t", [["구간", "실행"],
           ["17시 이후", "**0건**"],
           ["주말", "**0건**"],
           ["12시대", "공백 (점심)"]]),
    ("p", "설비는 무인으로도 돌 수 있으나 근무시간에만 가동됨. "
          "반송이 13~19초로 빠르고 가동률이 4~7%인데도 처리량이 낮은 이유. "
          "→ **작업자 인원·근무시간이 처리량을 결정**"),
])

# ── ③ 프린팅 — "왜 두 종류인가" 절 삽입 ─────────────────────────────────
insert_before("용도 (장비리스트 원문)", [
    ("h", "왜 두 종류인가 — 재료 점도가 방식을 가른다"),
    ("p", "같은 인쇄지만 다룰 수 있는 재료의 묽기가 **1만 배 차이**라 한 장비로 둘 다 못 함"),
    ("t", [["", "스크린", "잉크젯"],
           ["재료 점도", "**30,000 ~ 200,000 cps** (치약·꿀 수준)", "**1 ~ 20 cPs** (물 수준)"],
           ["재료를 내보내는 법", "제판 구멍으로 스퀴지가 밀어넣음", "노즐에서 액적으로 뿌림"],
           ["결과", "두껍게 · 거칠게", "얇게 · 정밀하게"],
           ["만드는 것", "전극 · 배선 · **격벽**", "**발광층** · 컬러필터 등 기능층"]]),
    ("p", "걸쭉한 페이스트는 노즐을 막아 잉크젯으로 못 쓰고, 묽은 잉크는 제판 구멍으로 흘러내려 스크린으로 못 씀. "
          "**만들려는 것이 두꺼운 구조물이냐 얇은 기능층이냐로 갈림**"),
    ("p", "컬러필터는 양쪽에 다 나오는데, **스크린이 격벽(칸막이)을 찍고 잉크젯이 그 안에 색을 채우는** 역할 분담으로 보임"),
], level=3)

doc.save(DST)
print(f"v2 복사 후 부분 교체 완료 → {DST}")
print(f"  표 {len(doc.tables)}개 · 문단 {len(doc.paragraphs)}개")
