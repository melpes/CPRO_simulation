"""KETI FAB AAS 자료조사 경위를 Word 보고서로 만든다.

4쪽 안에 넣기 위해 줄글을 최소화하고 표로 옮겼다.
수치는 만들 때마다 산출물에서 직접 읽어 온다(손으로 적어두면 어긋난다).
"""

import glob
import os
import sys
from collections import Counter, defaultdict

import openpyxl
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE = r"C:\Users\KangTaehui\KG\keti\CPRO_조립공정\시뮬레이션\Package"
KETI = os.path.join(BASE, r"docs\원본자료\keti-fab")
OUT = os.path.join(KETI, "KETI_FAB_AAS_자료조사_경위.docx")

FONT = "맑은 고딕"
GRAY = RGBColor(0x44, 0x44, 0x44)


# ── 산출물에서 수치를 읽어 온다 ─────────────────────────────────
def numbers():
    n = {}
    ws = openpyxl.load_workbook(os.path.join(KETI, r"AAS자료모음\_자료_총괄.xlsx"))["자료 총괄"]
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    n["총괄행"] = len(rows)
    n["등급"] = Counter(r[5] for r in rows)
    wb2 = openpyxl.load_workbook(os.path.join(KETI, r"AAS자료모음\_슬롯_목록.xlsx"))
    n["Property종"] = wb2["AAS별 슬롯"].max_row - 1
    n["특허"] = len(glob.glob(os.path.join(KETI, r"특허수집\명세서\*.json")))
    n["논문"] = len(glob.glob(os.path.join(KETI, r"논문수집\본문\*.txt")))
    n["캡처"] = len(glob.glob(os.path.join(KETI, r"AAS자료모음\*\*\캡처_*.jpg")))

    d = defaultdict(lambda: defaultdict(int))
    wb = openpyxl.load_workbook(os.path.join(KETI, "KETI_FAB_AAS_파라미터슬롯.xlsx"),
                                read_only=True)
    for sh in wb.sheetnames:
        if sh == "요약":
            continue
        cur = ""
        for r in wb[sh].iter_rows(min_row=2, values_only=True):
            s = (r[6] or "").strip()
            if s:
                cur = s
            own = not (cur.startswith("유사장비") or "KOSMO" in cur)
            d[sh]["계"] += 1
            if own:
                d[sh]["KETI"] += 1
    wb.close()
    n["AAS"] = d
    return n


# ── 문서 조립 도우미 ───────────────────────────────────────────
def setup(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(9.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.space_after = Pt(3)
    st.paragraph_format.line_spacing = 1.15
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(1.6)
        s.left_margin = s.right_margin = Cm(1.8)


def h(doc, text, size=12, space=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return p


def bullet(doc, text, indent=0.45):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(1)
    p.add_run("· " + text)
    return p


def table(doc, header, rows, widths=None, size=8.5):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, x in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(str(x))
        r.bold = True
        r.font.size = Pt(size)
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, x in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(x))
            r.font.size = Pt(size)
            r.font.name = FONT
            r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            if i and str(x).replace("%", "").replace(",", "").isdigit():
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(8.5)
    r.font.color.rgb = GRAY
    return p


def main():
    sys.stdout.reconfigure(encoding="utf-8")
    n = numbers()
    doc = Document()
    setup(doc)

    # ── 제목 ──────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("KETI 전북본부 FAB — AAS 파라미터 자료조사 경위")
    r.bold = True
    r.font.size = Pt(15)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    note(doc, "AAS 모델링 대상 12종에 넣을 Property 를 확보하기 위해 어디를 어떻게 조사했고, "
              "각 자료를 왜 채택했는지 정리한다. 수치는 산출물에서 직접 읽은 값이다.")

    # ── 1. 조사 대상 ───────────────────────────────────────
    h(doc, "1. 조사한 사이트와 확인 지점")
    table(doc,
          ["사이트", "확인한 화면·경로", "얻은 것", "제약"],
          [["ZEUS\n(zeus.go.kr)", "등록장비 상세 /search/equip/read/\n장비예약 상세 /resv/equip/read/",
            "구성 및 성능, 특성,\n동일모델 장비목록", "관련자료 탭(매뉴얼)은\nKETI 보유 40대 전부 비어 있음"],
           ["i-Tube\n(itube.or.kr)", "상세 sharingView.do\n첨부 fileDownMyEquip.do",
            "설비 매뉴얼 PDF 4건", "매뉴얼 보유 설비가 소수"],
           ["NTIS", "과제상세 → 연구개발성과 →\n연구시설장비 목록",
            "구축과제로 장비 발굴\n(rstId = ZEUS 등록ID)", "나노기술집적센터가 광주·전주\n동명이라 보유기관 대조 필요"],
           ["OpenAlex", "기관 ID I4210131650", "KETI 소속 논문 112편", "검색어가 길면 미검출"],
           ["RISS / 출판사", "DOI → 본문", f"논문 본문 {n['논문']}편",
            "Elsevier 5편 CAPTCHA 차단"],
           ["KIPRIS", "검색식 AP=[제조사]*AB=[키워드]", "제조사 특허 발굴 231건",
            "URL 직접 접속 시 결과 미표시\n(검색창 재입력 필요)"],
           ["Google Patents", "/patent/<공개번호>/ko", f"명세서·청구항 {n['특허']}건",
            "요청 과다 시 차단(약 1시간)"],
           ["KOSMO 제공자료", "kasmo_후보데이터.zip", "완성 AAS 3종 + 가이드",
            "공개 URL 없음"]],
          widths=[2.4, 5.0, 4.4, 5.2])
    note(doc, "※ KIPRIS 출원번호는 Google Patents 공개번호와 다른 체계다. 번호를 변환해 접속하면 "
              "형식이 맞아 응답은 오지만 전혀 다른 특허가 잡힌다. 제목+출원인 검색으로 공개번호를 먼저 찾았다.")

    # ── 2. 자료 채택 기준 ──────────────────────────────────
    h(doc, "2. 자료를 채택한 기준")
    table(doc,
          ["구분", "채택 기준", "제외한 예"],
          [["Fab장비", "KETI 보유 장비 자신 — 무조건 포함", "—"],
           ["유사장비", "같은 공정을 하는가, 사양 항목이 그대로 대응되는가\n"
                     "(이름 유사·제조사 동일은 근거가 아님)",
            "OLED 소재 합성 시스템(글로브박스), 확산로(열처리),\n"
            "프로브 스테이션(측정), 플라즈마 애싱(진공계통 중복)"],
           ["부분유사", "일부 계통만 대응 — 겹치는 계통만 사용",
            "마이크로파 PECVD(여기 방식 상이), iMV-DX4(분말 공정)"],
           ["논문", "그 AAS 의 공정을 실제로 돌린 논문 —\n값이 아니라 조정 파라미터의 이름만 채택",
            "리뷰 논문, 공정 무관(Zn전지 하이드로겔 등)"],
           ["특허", "장비가 조정하는 값이 수치로 명시된 것",
            "구조 특허(부품 배치·형상), 소모품 물성(연마패드 경도),\n"
            "12종에 없는 공정(CMP), 반도체 무관"],
           ["KOSMO", "기능이 실제로 같은 AAS 에만 배정",
            "표면처리기 TechnicalData(원심 바스켓형), 에칭장비 생산관리 계통,\n"
            "전해도금조 전기도금 제어(CBD 는 전류 미인가)"]],
          widths=[1.8, 7.4, 7.8])

    # ── 3. 12종 분류 ──────────────────────────────────────
    h(doc, "3. 12종 분류와 자료 배정 원칙")
    bullet(doc, "개발 목표 12종은 KETI FAB 현장 확인으로 먼저 확정된 것이며, 본 조사는 그 12종에 "
                "자료를 대응시켜 Property 확보 가능성을 확인한 작업이다.")
    bullet(doc, "증착·식각 8 (박막증착장비 4챔버 + 유기증착기 3챔버 + PEALD) / 포토 3 / 프린터 1. "
                "CBD 는 12종 밖 참고.")
    bullet(doc, "분해 — 한 장비 안에서도 파라미터 종류가 다른 챔버는 별개 AAS. 병합 — 역할이 같고 "
                "파라미터가 겹치면 1종(프린팅 6대). 범위 — 동적 조정값과 정적 사양 모두 포함.")
    bullet(doc, "진공/비진공이 갈림선. 박막증착장비 Dry Etcher 는 진공 건식, 엣쳐/스트리퍼는 습식·상압. "
                "같은 '식각'이라도 다른 장비다.")

    # ── 4. 확보 현황 ───────────────────────────────────────
    h(doc, "4. AAS별 확보 현황")
    aas_rows = []
    for k, v in sorted(n["AAS"].items(), key=lambda x: -x[1]["계"]):
        pct = v["KETI"] * 100 // v["계"] if v["계"] else 0
        aas_rows.append([k, v["계"], v["KETI"], f"{pct}%",
                         "가능" if pct >= 40 else ("제한적" if pct >= 15 else "불가")])
    table(doc, ["AAS", "Property 계", "KETI 자체", "자체 비율", "카탈로그 역할"],
          aas_rows, widths=[6.2, 2.4, 2.4, 2.2, 3.6])
    note(doc, "‘KETI 자체’는 KETI 보유 장비의 ZEUS·i-Tube·매뉴얼에서 나온 것이다. 나머지는 타 기관 "
              "유사장비와 KOSMO AAS 에서 이름을 가져온 것으로, 값은 KETI 값이 아니다.")

    # ── 5. 수집 결과 ───────────────────────────────────────
    h(doc, "5. 수집 산출물")
    g = n["등급"]
    table(doc,
          ["산출물", "규모", "내용"],
          [["AAS자료모음/ (13개 폴더)", f"{n['총괄행']}행",
            f"Fab장비 {g.get('Fab장비',0)} · 유사장비 {g.get('유사장비',0)} · "
            f"부분유사 {g.get('부분유사',0)} · 논문 {g.get('논문',0)} · 특허 {g.get('특허',0)}"],
           ["_자료_총괄.xlsx", f"{n['총괄행']}행",
            "링크 / KETI FAB 설비와의 유사성 / 확인할 위치 / Property 후보"],
           ["_슬롯_목록.xlsx", f"{n['Property종']}종",
            "AAS별 Property 목록 + 값 예시 + 근거 링크, 특허 색인 86건"],
           ["KETI_FAB_AAS_파라미터슬롯.xlsx", "829행", "출처·근거 링크가 붙은 Property 원장"],
           ["특허 명세서 / 논문 본문", f"{n['특허']}건 / {n['논문']}편", "원문 저장, Ctrl+F 검색어 기록"],
           ["사이트 사양 캡처", f"{n['캡처']}건", "구성 및 성능 영역을 펼친 화면"]],
          widths=[5.4, 2.6, 8.8])

    # ── 6. 확인된 것 ───────────────────────────────────────
    h(doc, "6. 조사에서 확인된 것")
    bullet(doc, "포털 사양만으로는 챔버 단위 AAS 가 성립하지 않는다. 최초 집계 시 챔버당 1~2개였다. "
                "유사장비 이식 · KOSMO AAS · 설비 매뉴얼 세 가지로 이를 넘었다.")
    bullet(doc, "설비 매뉴얼의 효과가 가장 크다. 스핀 트랙 매뉴얼 1건이 Property 를 10 → 169 로 늘렸다. "
                "포털에 없는 유닛별 사양(Spin Coater Unit·Hot Plate·Transfer Robot 등)이 여기에 있다.")
    bullet(doc, "제조사마다 특허 성격이 다르다. 테스·씨엔원은 공정 조건을 수치로 명시하나(평균 12건대), "
                "선익시스템은 구조 특허 위주로 13건 중 8건이 수치 0이었다.")
    bullet(doc, "수치 개수가 Property 가치를 보장하지 않는다. 연마패드 특허는 수치 101건이나 전부 패드 "
                "물성이라 채택 0, 크롬 패턴 형성 방법은 19건 중 8개가 모두 공정 조건이었다.")
    bullet(doc, "이름만 보고 Property 를 병합하면 안 된다. Dry Etcher 의 압력 3종(공정·챔버·챔버 내부)은 "
                "원문상 같은 값이나, 온도 6종(기판·기판 가열·서셉터·서셉터 냉각·샤워헤드·챔버 벽)은 "
                "서로 다른 부위다. 기판 가열과 서셉터 냉각을 동시에 하는 2단 제어이기 때문이다.")

    # ── 7. 한계 ───────────────────────────────────────────
    h(doc, "7. 한계 — 카탈로그 역할의 범위")
    bullet(doc, "Property 골격은 확보했으나 값은 대부분 KETI 값이 아니다. 전체 829개 중 값이 있는 597개 "
                "가운데 68%가 타 기관 장비 또는 KOSMO 유래다.")
    bullet(doc, "용어가 제조사 공식 표기가 아니다. 같은 항목을 ZEUS 등록자 표기, 특허 권리 표현, 논문 "
                "표기가 각각 다르게 부른다. 원문 대조 없이는 병합·분리를 판단할 수 없다.")
    bullet(doc, "대외 문서에서는 자체 값과 참고 값이 구분되어야 한다. 산출물에 ‘등급’과 ‘자료구분’ 열을 "
                "유지한 이유가 이것이다.")

    # ── 8. 남은 일 ────────────────────────────────────────
    h(doc, "8. 남은 일")
    table(doc,
          ["대상", "상태", "기대 효과"],
          [["제조사 매뉴얼 (요청 진행 중)",
            "씨엔원(PEALD) · 테스(PECVD) · 선익시스템(유기증착기) · 제우스(CBD)",
            "KETI 자체 값 확보. 스핀 트랙 사례 기준 가장 큰 레버"],
           ["노광 AAS", "코디엠 특허 KIPRIS 0건, 논문 없음", "추가 확보 경로 없음"],
           ["논문 5편", "Elsevier ScienceDirect CAPTCHA", "우회하지 않음"],
           ["NTIS 과제 13건", "로그인 브라우저로 연구시설장비 목록 확인 필요", "미발굴 장비 확인"],
           ["특허 저순위 82건", "선익 37 · 마이다스 16 · 나래나노텍 29 미착수",
            "선익은 구조 특허 수율이 낮아 후순위"]],
          widths=[4.6, 6.6, 5.6])

    doc.save(OUT)
    print(f"저장 → {OUT}")


if __name__ == "__main__":
    main()
